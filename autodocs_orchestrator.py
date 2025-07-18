import os
import datetime
import threading
import time
from pathlib import Path
from typing import List, Dict, Optional, Callable
import json

# Import our existing modules
from audiovisual.av_trigger import record as record_clip
from transcribe.transcribe_summary import transcribe_audio, summarize_transcription


class AutoDocsOrchestrator:
    """
    Main orchestrator class that manages recording clips and generating documentation.
    
    This class coordinates:
    1. Recording audio/video clips
    2. Transcribing audio to text
    3. Summarizing transcriptions
    4. Organizing clips into a structured document
    """
    
    def __init__(self, output_dir: str = "autodocs_output"):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)
        
        self.clips: List[Dict] = []
        self.session_id = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        self.session_dir = self.output_dir / f"session_{self.session_id}"
        
        # Create session directory and subdirectories
        self.session_dir.mkdir(parents=True, exist_ok=True)
        
        # Create subdirectories for organization
        (self.session_dir / "clips").mkdir(exist_ok=True)
        (self.session_dir / "audio").mkdir(exist_ok=True)
        (self.session_dir / "gifs").mkdir(exist_ok=True)
        (self.session_dir / "transcripts").mkdir(exist_ok=True)
        
        self.status_callback: Optional[Callable] = None
        
    def set_status_callback(self, callback: Callable[[str], None]):
        """Set a callback function to receive status updates"""
        self.status_callback = callback
        
    def _update_status(self, message: str):
        """Internal method to update status"""
        print(f"[AutoDocs] {message}")
        if self.status_callback:
            self.status_callback(message)
    
    def record_clip(self, duration: int = 15, title: str = None) -> Dict:
        """
        Record a new clip with audio and video
        
        Args:
            duration: Recording duration in seconds
            title: Optional title for the clip
            
        Returns:
            Dict containing clip information
        """
        if title is None:
            title = f"Clip {len(self.clips) + 1}"
            
        self._update_status(f"🔴 Recording clip: {title}")
        
        # Create a timestamp for this clip
        clip_timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        
        # Change to the clips directory for recording
        original_dir = os.getcwd()
        clips_dir = self.session_dir / "clips"
        
        # Ensure the clips directory exists
        clips_dir.mkdir(parents=True, exist_ok=True)
        os.chdir(clips_dir)
        
        try:
            # Record the clip using existing av_trigger functionality
            record_clip(duration=duration, status_callback=self._update_status)
            
            # Find the generated files (they'll be the most recent ones)
            audio_files = list(Path(".").glob("audio_*.wav"))
            gif_files = list(Path(".").glob("screen_*.gif"))
            video_files = list(Path(".").glob("screen_*.mp4"))
            
            if not audio_files or not gif_files:
                raise Exception("Recording files not found")
                
            # Get the most recent files
            audio_file = max(audio_files, key=os.path.getctime)
            gif_file = max(gif_files, key=os.path.getctime)
            video_file = max(video_files, key=os.path.getctime) if video_files else None
            
            # Move files to absolute paths for storage
            audio_file = Path(audio_file).resolve()
            gif_file = Path(gif_file).resolve()
            video_file = Path(video_file).resolve() if video_file else None
            
            # Create clip metadata
            clip_data = {
                "id": len(self.clips) + 1,
                "title": title,
                "timestamp": clip_timestamp,
                "duration": duration,
                "audio_file": str(audio_file),
                "gif_file": str(gif_file),
                "video_file": str(video_file) if video_file else None,
                "transcription": None,
                "summary": None,
                "status": "recorded"
            }
            
            self.clips.append(clip_data)
            self._save_session_metadata()
            
            self._update_status(f"✅ Clip recorded: {title}")
            return clip_data
            
        finally:
            os.chdir(original_dir)
    
    def process_clip(self, clip_id: int) -> Dict:
        """
        Process a recorded clip by transcribing and summarizing
        
        Args:
            clip_id: ID of the clip to process
            
        Returns:
            Updated clip data
        """
        clip = self._get_clip_by_id(clip_id)
        if not clip:
            raise ValueError(f"Clip with ID {clip_id} not found")
            
        self._update_status(f"🔄 Processing clip: {clip['title']}")
        
        try:
            # Transcribe audio
            self._update_status(f"🎵 Transcribing audio for: {clip['title']}")
            transcription = transcribe_audio(clip['audio_file'])
            clip['transcription'] = transcription
            
            # Save transcription to file
            transcript_file = self.session_dir / "transcripts" / f"clip_{clip_id}_transcript.txt"
            with open(transcript_file, 'w', encoding='utf-8') as f:
                f.write(transcription)
            clip['transcript_file'] = str(transcript_file)
            
            # Generate summary
            self._update_status(f"📝 Generating summary for: {clip['title']}")
            summary = summarize_transcription(transcription)
            clip['summary'] = summary
            
            # Save summary to file
            summary_file = self.session_dir / "transcripts" / f"clip_{clip_id}_summary.txt"
            with open(summary_file, 'w', encoding='utf-8') as f:
                f.write(summary)
            clip['summary_file'] = str(summary_file)
            
            clip['status'] = 'processed'
            self._save_session_metadata()
            
            self._update_status(f"✅ Clip processed: {clip['title']}")
            return clip
            
        except Exception as e:
            self._update_status(f"❌ Error processing clip {clip['title']}: {str(e)}")
            clip['status'] = 'error'
            clip['error'] = str(e)
            self._save_session_metadata()
            raise
    
    def process_all_clips(self):
        """Process all recorded clips that haven't been processed yet"""
        unprocessed_clips = [clip for clip in self.clips if clip['status'] == 'recorded']
        
        if not unprocessed_clips:
            self._update_status("No clips to process")
            return
            
        self._update_status(f"📋 Processing {len(unprocessed_clips)} clips...")
        
        for clip in unprocessed_clips:
            try:
                self.process_clip(clip['id'])
            except Exception as e:
                self._update_status(f"❌ Failed to process clip {clip['id']}: {str(e)}")
                continue
    
    def generate_word_document(self) -> str:
        """
        Generate a Word document with all clips, summaries, and GIFs
        
        Returns:
            Path to the generated Word document
        """
        try:
            from docx import Document
            from docx.shared import Inches
        except ImportError:
            raise ImportError("python-docx is required. Install it with: pip install python-docx")
        
        self._update_status("📄 Generating Word document...")
        
        doc = Document()
        
        # Add title
        title = doc.add_heading(f'AutoDocs Tutorial - {self.session_id}', 0)
        
        # Add metadata
        doc.add_paragraph(f"Generated on: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"Total Steps: {len(self.clips)}")
        doc.add_paragraph("")  # Empty line
        
        # Add each clip as a step
        for i, clip in enumerate(self.clips, 1):
            # Step heading
            doc.add_heading(f'Step {i}: {clip["title"]}', level=1)
            
            # Add summary if available
            if clip.get('summary'):
                doc.add_heading('Summary', level=2)
                doc.add_paragraph(clip['summary'])
            
            # Add GIF if available
            if clip['gif_file'] and os.path.exists(clip['gif_file']):
                doc.add_heading('Screen Recording', level=2)
                try:
                    # Copy GIF to session directory for better organization
                    gif_filename = f"step_{i}_recording.gif"
                    gif_dest = self.session_dir / "gifs" / gif_filename
                    
                    import shutil
                    shutil.copy2(clip['gif_file'], gif_dest)
                    
                    # Convert GIF to static image for Word document
                    # Since Word doesn't support animated GIFs, we'll use the first frame
                    from PIL import Image
                    static_image_path = gif_dest.with_suffix('.png')
                    
                    # Open GIF and save the gif
                    run.add_picture(str(gif_dest), width=Inches(6))
                    
                    # Add the gif to the Word document
                    paragraph = doc.add_paragraph()
                    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                    
                    # Add image with reasonable size (max width of 6 inches)
                    try:
                        run.add_picture(str(static_image_path), width=Inches(6))
                    except Exception as img_error:
                        # If image insertion fails, fall back to text reference
                        doc.add_paragraph(f"[Screen Recording: {gif_filename}]")
                        doc.add_paragraph(f"Note: Could not embed image - {str(img_error)}")
                    
                    # Add note about animated version
                    doc.add_paragraph(f"Animated GIF available at: {gif_dest}", style='Caption')
                    
                except Exception as e:
                    doc.add_paragraph(f"[Error processing screen recording: {str(e)}]")
            
            # Add full transcription in a collapsible-style format
            if clip.get('transcription'):
                doc.add_heading('Full Transcription', level=2)
                transcription_para = doc.add_paragraph()
                transcription_para.add_run(clip['transcription'])
                transcription_para.style = 'Quote'
            
            # Add metadata
            doc.add_heading('Technical Details', level=2)
            details = [
                f"Duration: {clip['duration']} seconds",
                f"Recorded: {clip['timestamp']}",
                f"Status: {clip['status']}"
            ]
            for detail in details:
                doc.add_paragraph(detail, style='List Bullet')
            
            # Add page break except for last clip
            if i < len(self.clips):
                doc.add_page_break()
        
        # Save document
        doc_path = self.session_dir / f"AutoDocs_Tutorial_{self.session_id}.docx"
        doc.save(doc_path)
        
        self._update_status(f"✅ Word document generated: {doc_path}")
        return str(doc_path)
    
    def generate_markdown_document(self) -> str:
        """
        Generate a Markdown document with all clips, summaries, and GIFs
        
        Returns:
            Path to the generated Markdown document
        """
        self._update_status("📄 Generating Markdown document...")
        
        doc_path = self.session_dir / f"AutoDocs_Tutorial_{self.session_id}.md"
        
        with open(doc_path, 'w', encoding='utf-8') as f:
            # Header
            f.write(f"# AutoDocs Tutorial - {self.session_id}\n\n")
            f.write(f"Generated on: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")
            f.write(f"Total Steps: {len(self.clips)}\n\n")
            f.write("---\n\n")
            
            # Table of Contents
            f.write("## Table of Contents\n\n")
            for i, clip in enumerate(self.clips, 1):
                f.write(f"- [Step {i}: {clip['title']}](#step-{i}-{clip['title'].lower().replace(' ', '-')})\n")
            f.write("\n---\n\n")
            
            # Steps
            for i, clip in enumerate(self.clips, 1):
                f.write(f"## Step {i}: {clip['title']}\n\n")
                f.write(f"**Duration:** {clip['duration']} seconds\n\n")
                f.write(f"**Recorded:** {clip['timestamp']}\n\n")
                
                # Summary
                if clip.get('summary'):
                    f.write(f"### Summary\n\n")
                    f.write(f"{clip['summary']}\n\n")
                
                # GIF
                if clip['gif_file'] and os.path.exists(clip['gif_file']):
                    # Copy GIF to gifs directory
                    gif_filename = f"step_{i}_recording.gif"
                    gif_dest = self.session_dir / "gifs" / gif_filename
                    
                    import shutil
                    shutil.copy2(clip['gif_file'], gif_dest)
                    
                    f.write(f"### Screen Recording\n\n")
                    f.write(f"![Step {i} Recording](gifs/{gif_filename})\n\n")
                    
                    # Add alternative formats for better compatibility
                    try:
                        # Also create a PNG version for static viewing
                        from PIL import Image
                        static_image_path = gif_dest.with_suffix('.png')
                        
                        with Image.open(gif_dest) as img:
                            if img.mode != 'RGB':
                                img = img.convert('RGB')
                            img.save(static_image_path, 'PNG')
                        
                        f.write(f"*Static version: [step_{i}_recording.png](gifs/step_{i}_recording.png)*\n\n")
                    except Exception:
                        pass  # If PNG conversion fails, just continue with GIF
                
                # Full transcription
                if clip.get('transcription'):
                    f.write(f"### Full Transcription\n\n")
                    f.write(f"<details>\n")
                    f.write(f"<summary>Click to expand full transcription</summary>\n\n")
                    f.write(f"```\n{clip['transcription']}\n```\n\n")
                    f.write(f"</details>\n\n")
                
                f.write("---\n\n")
        
        self._update_status(f"✅ Markdown document generated: {doc_path}")
        return str(doc_path)
    
    def generate_html_document(self) -> str:
        """
        Generate an HTML document with all clips, summaries, and GIFs
        HTML has better support for animated GIFs than Word documents
        
        Returns:
            Path to the generated HTML document
        """
        self._update_status("📄 Generating HTML document...")
        
        doc_path = self.session_dir / f"AutoDocs_Tutorial_{self.session_id}.html"
        
        with open(doc_path, 'w', encoding='utf-8') as f:
            # HTML Header
            f.write(f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>AutoDocs Tutorial - {self.session_id}</title>
    <style>
        body {{
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, Oxygen, Ubuntu, Cantarell, sans-serif;
            line-height: 1.6;
            max-width: 1200px;
            margin: 0 auto;
            padding: 20px;
            background-color: #f8f9fa;
        }}
        .container {{
            background: white;
            padding: 30px;
            border-radius: 8px;
            box-shadow: 0 2px 10px rgba(0,0,0,0.1);
        }}
        h1 {{
            color: #2c3e50;
            border-bottom: 3px solid #3498db;
            padding-bottom: 10px;
        }}
        h2 {{
            color: #34495e;
            margin-top: 30px;
        }}
        h3 {{
            color: #7f8c8d;
        }}
        .step {{
            margin: 40px 0;
            padding: 20px;
            border-left: 4px solid #3498db;
            background-color: #f8f9fa;
        }}
        .gif-container {{
            text-align: center;
            margin: 20px 0;
            padding: 15px;
            background: white;
            border-radius: 8px;
            box-shadow: 0 2px 5px rgba(0,0,0,0.1);
        }}
        .gif-container img {{
            max-width: 100%;
            border-radius: 4px;
            box-shadow: 0 4px 8px rgba(0,0,0,0.15);
        }}
        .summary {{
            background: #e8f4fd;
            padding: 15px;
            border-radius: 6px;
            border-left: 4px solid #3498db;
        }}
        .transcription {{
            background: #f8f9fa;
            padding: 15px;
            border-radius: 6px;
            font-family: 'Courier New', monospace;
            font-size: 14px;
        }}
        .metadata {{
            color: #7f8c8d;
            font-size: 14px;
        }}
        details {{
            margin: 15px 0;
        }}
        summary {{
            cursor: pointer;
            font-weight: bold;
            color: #2c3e50;
        }}
        .toc {{
            background: #ecf0f1;
            padding: 20px;
            border-radius: 6px;
            margin: 20px 0;
        }}
        .toc ul {{
            list-style-type: none;
            padding: 0;
        }}
        .toc li {{
            margin: 8px 0;
        }}
        .toc a {{
            text-decoration: none;
            color: #2980b9;
        }}
        .toc a:hover {{
            text-decoration: underline;
        }}
    </style>
</head>
<body>
    <div class="container">
        <h1>AutoDocs Tutorial - {self.session_id}</h1>
        <div class="metadata">
            <p><strong>Generated on:</strong> {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>
            <p><strong>Total Steps:</strong> {len(self.clips)}</p>
        </div>
        
        <div class="toc">
            <h2>Table of Contents</h2>
            <ul>
""")
            
            # Table of Contents
            for i, clip in enumerate(self.clips, 1):
                f.write(f'                <li><a href="#step-{i}">Step {i}: {clip["title"]}</a></li>\n')
            
            f.write("""            </ul>
        </div>
        
""")
            
            # Steps
            for i, clip in enumerate(self.clips, 1):
                f.write(f"""        <div class="step" id="step-{i}">
            <h2>Step {i}: {clip['title']}</h2>
            <div class="metadata">
                <strong>Duration:</strong> {clip['duration']} seconds | 
                <strong>Recorded:</strong> {clip['timestamp']}
            </div>
            
""")
                
                # Summary
                if clip.get('summary'):
                    f.write(f"""            <h3>Summary</h3>
            <div class="summary">
                {clip['summary']}
            </div>
            
""")
                
                # GIF
                if clip['gif_file'] and os.path.exists(clip['gif_file']):
                    # Copy GIF to gifs directory
                    gif_filename = f"step_{i}_recording.gif"
                    gif_dest = self.session_dir / "gifs" / gif_filename
                    
                    import shutil
                    shutil.copy2(clip['gif_file'], gif_dest)
                    
                    f.write(f"""            <h3>Screen Recording</h3>
            <div class="gif-container">
                <img src="gifs/{gif_filename}" alt="Step {i} Recording" loading="lazy">
                <p><em>Animated demonstration of Step {i}</em></p>
            </div>
            
""")
                
                # Full transcription
                if clip.get('transcription'):
                    f.write(f"""            <details>
                <summary>Full Transcription</summary>
                <div class="transcription">
                    {clip['transcription'].replace(chr(10), '<br>')}
                </div>
            </details>
            
""")
                
                f.write("        </div>\n\n")
            
            f.write("""    </div>
</body>
</html>""")
        
        self._update_status(f"✅ HTML document generated: {doc_path}")
        return str(doc_path)
    
    def _get_clip_by_id(self, clip_id: int) -> Optional[Dict]:
        """Get a clip by its ID"""
        for clip in self.clips:
            if clip['id'] == clip_id:
                return clip
        return None
    
    def _save_session_metadata(self):
        """Save session metadata to JSON file"""
        # Ensure the session directory exists before trying to save
        self.session_dir.mkdir(parents=True, exist_ok=True)
        
        metadata_file = self.session_dir / "session_metadata.json"
        try:
            with open(metadata_file, 'w', encoding='utf-8') as f:
                json.dump({
                    'session_id': self.session_id,
                    'created': datetime.datetime.now().isoformat(),
                    'clips': self.clips
                }, f, indent=2)
        except Exception as e:
            print(f"Warning: Could not save session metadata: {e}")
            # Don't raise the exception, just log it
            pass
    
    def load_session(self, session_dir: str):
        """Load an existing session from directory"""
        session_path = Path(session_dir)
        metadata_file = session_path / "session_metadata.json"
        
        if not metadata_file.exists():
            raise ValueError(f"No session metadata found in {session_dir}")
        
        with open(metadata_file, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        self.session_id = data['session_id']
        self.clips = data['clips']
        self.session_dir = session_path
        
        self._update_status(f"📂 Loaded session: {self.session_id}")
    
    def get_session_summary(self) -> Dict:
        """Get a summary of the current session"""
        processed_count = len([c for c in self.clips if c['status'] == 'processed'])
        error_count = len([c for c in self.clips if c['status'] == 'error'])
        
        return {
            'session_id': self.session_id,
            'total_clips': len(self.clips),
            'processed_clips': processed_count,
            'error_clips': error_count,
            'session_dir': str(self.session_dir)
        }


# Convenience functions for quick operations
def quick_record_and_process(duration: int = 15, title: str = None) -> str:
    """
    Quick function to record a single clip and generate a document
    
    Args:
        duration: Recording duration in seconds
        title: Optional title for the clip
        
    Returns:
        Path to the generated Word document
    """
    orchestrator = AutoDocsOrchestrator()
    
    # Record clip
    clip = orchestrator.record_clip(duration=duration, title=title)
    
    # Process clip
    orchestrator.process_clip(clip['id'])
    
    # Generate document
    doc_path = orchestrator.generate_word_document()
    
    return doc_path


def create_multi_clip_session(clip_configs: List[Dict]) -> str:
    """
    Create a session with multiple clips
    
    Args:
        clip_configs: List of dictionaries with 'duration' and 'title' keys
        
    Returns:
        Path to the generated document
    """
    orchestrator = AutoDocsOrchestrator()
    
    # Record all clips
    for config in clip_configs:
        duration = config.get('duration', 15)
        title = config.get('title')
        orchestrator.record_clip(duration=duration, title=title)
    
    # Process all clips
    orchestrator.process_all_clips()
    
    # Generate document
    doc_path = orchestrator.generate_word_document()
    
    return doc_path
